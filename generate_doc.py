from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x0D, 0x1E, 0x3A)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x0D, 0x1E, 0x3A)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0xF2, 0x65, 0x22)
    return p

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

def label_val(label, val):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(val)

def divider():
    doc.add_paragraph('─' * 60)

# ── COVER ──────────────────────────────────────────────────────────
title = doc.add_heading('AGENTIX', 0)
title.runs[0].font.color.rgb = RGBColor(0xF2, 0x65, 0x22)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Complete Website Content — All Pages')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True

meta = doc.add_paragraph('AI Tools Marketplace  |  ai-agentix.com  |  May 2026')
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── HOMEPAGE ───────────────────────────────────────────────────────
h1('HOMEPAGE')

h2('Hero Section')
label_val('Headline', 'The AI Stack for Modern Business')
label_val('Subheadline', 'Content. Sales. Research. Operations. One platform — infinite scale.')
label_val('CTA 1', 'Explore Tools → /category/content')
label_val('CTA 2', 'Watch Demo')
label_val('Stats', '16 AI Tools  |  4 Business Categories  |  10× Faster')

h2('Categories Section')
cats = [
    ('Content AI', '#7C3AED', 'Create anything. Publish everywhere.', '10× Faster content creation', '/category/content'),
    ('Sales AI', '#F26522', 'Find leads. Close deals. Automatically.', '3× More pipeline, same team', '/category/sales'),
    ('Research AI', '#0EA5E9', 'Know your market before your competitors do.', '48h To full market analysis', '/category/market-research'),
    ('Business AI', '#059669', 'Automate operations. Scale without limits.', '60% Reduction in operational cost', '/category/business'),
]
for name, color, tagline, stat, link in cats:
    h3(name)
    label_val('Tagline', tagline)
    label_val('Stat', stat)
    label_val('Link', link)

h2('Featured Tools')
for tool in ['AI Content Generator', 'Lead Scraper', 'Competitor Analyzer', 'CRM System']:
    bullet(tool)

h2('How It Works')
for step in ['1. Pick your category', '2. Choose your AI tools', '3. Go live in minutes']:
    bullet(step)

h2('Trust Stats')
for stat in ['16 Tools', '4 Categories', '99.9% Uptime', '24/7 Support']:
    bullet(stat)

doc.add_page_break()

# ── CATEGORIES ─────────────────────────────────────────────────────
h1('CATEGORY PAGES  (/category/:slug)')

category_data = [
    {
        'name': 'Content AI', 'slug': 'content',
        'tagline': 'Create anything. Publish everywhere.',
        'description': 'AI-powered tools that write, schedule, edit, and automate your entire content pipeline — from first draft to final post.',
        'stat': '10× Faster content creation',
        'tools': ['AI Content Generator', 'Social Media Scheduler', 'AI Photo Editor', 'Workflow Generator'],
    },
    {
        'name': 'Sales AI', 'slug': 'sales',
        'tagline': 'Find leads. Close deals. Automatically.',
        'description': 'Intelligent sales automation tools that scrape leads, automate outreach, qualify prospects via voice, and keep your CRM perfectly synced.',
        'stat': '3× More pipeline, same team',
        'tools': ['Lead Scraper', 'WhatsApp Automation', 'Cold Calling Bot', 'CRM Sync'],
    },
    {
        'name': 'Research AI', 'slug': 'market-research',
        'tagline': 'Know your market before your competitors do.',
        'description': 'Real-time competitive intelligence, dynamic pricing analysis, and AI-generated go-to-market strategies based on live market data.',
        'stat': '48h To full market analysis',
        'tools': ['Competitor Analyzer', 'Pricing Intelligence', 'Market Strategy Generator'],
    },
    {
        'name': 'Business AI', 'slug': 'business',
        'tagline': 'Automate operations. Scale without limits.',
        'description': 'Enterprise-grade CRM, ERP, LMS, website builder, and e-commerce tools with AI embedded at every layer — so your business runs smarter.',
        'stat': '60% Reduction in operational cost',
        'tools': ['CRM System', 'ERP Modules', 'LMS System', 'Website Builder', 'E-Commerce Sites'],
    },
]

for cat in category_data:
    h2(f"{cat['name']} — /category/{cat['slug']}")
    label_val('Tagline', cat['tagline'])
    label_val('Description', cat['description'])
    label_val('Key Stat', cat['stat'])
    p = doc.add_paragraph()
    p.add_run('Tools: ').bold = True
    p.add_run(', '.join(cat['tools']))

doc.add_page_break()

# ── TOOL PAGES ─────────────────────────────────────────────────────
h1('TOOL PAGES  (/tools/:slug)')

tools_data = [
    # CONTENT
    {
        'name': 'AI Content Generator', 'slug': 'ai-content-generator', 'category': 'Content AI',
        'tagline': 'From brief to published — in seconds.',
        'description': 'Generate long-form blog posts, ad copy, email sequences, product descriptions, and social content at scale. Our AI learns your brand voice and produces content indistinguishable from your best writer — in a fraction of the time.',
        'stats': ['10× Faster than manual writing', '94% Brand voice accuracy', '60+ Content formats supported'],
        'features': ['Brand Voice Engine', 'Multi-Format Output', 'SEO Optimisation', 'Batch Generation'],
        'how_it_works': ['Define your brief', 'AI generates drafts', 'Review and publish'],
        'ai_role': 'Researches the topic, structures the outline, writes all drafts, checks SEO, flags readability issues, and formats output for your target platform.',
        'human_role': 'Provides the brief, reviews tone accuracy, approves final copy, and applies strategic editorial judgement.',
        'use_cases': [
            'Marketing Agency — 12× output increase, same team size',
            'E-Commerce Brand — $180K saved vs copywriting agency',
            'SaaS Company — 430% increase in organic traffic in 6 months',
        ],
        'integrations': 'WordPress, HubSpot, Webflow, Notion, Google Docs, Shopify',
        'faqs': [
            ('How does it learn my brand voice?', 'Upload 5–10 examples of your best-performing content. The model fine-tunes to your style and applies it to every future output.'),
            ('Can it write in multiple languages?', 'Yes. We support 40+ languages with native-quality output for English, Spanish, French, German, Hindi, Arabic, and more.'),
            ('Is the content plagiarism-free?', 'All content is generated from scratch. We run a built-in plagiarism check before delivery to confirm 0% matches.'),
            ('What is the maximum article length?', 'Up to 10,000 words per generation. Long-form whitepapers, guides, and reports are fully supported.'),
            ('Does it integrate with my CMS?', 'Direct integrations with WordPress, HubSpot, Webflow, and Notion. Any other CMS can be connected via our API.'),
        ],
    },
    {
        'name': 'Social Media Scheduler', 'slug': 'social-media-scheduler', 'category': 'Content AI',
        'tagline': 'Schedule smarter. Grow faster.',
        'description': 'AI-powered scheduling that analyses your audience behaviour, identifies peak engagement windows, writes captions, and posts across all platforms without manual effort.',
        'stats': ['4.2× Average engagement lift', '12 Platforms connected', '30 min Weekly time investment'],
        'features': ['Optimal Timing AI', 'Caption Generator', 'Content Calendar', 'Analytics Dashboard'],
        'how_it_works': ['Connect your accounts', 'Upload or create content', 'AI publishes and optimises'],
        'ai_role': 'Analyses audience behaviour, writes captions, selects optimal posting times, formats content per platform, and adjusts strategy based on performance data.',
        'human_role': 'Creates or approves visual assets, sets campaign objectives, reviews monthly performance reports.',
        'use_cases': [
            'Personal Brand — 3.8× follower growth in 90 days',
            'Retail Brand — 42% increase in social-driven foot traffic',
            'Content Agency — scaled 4× without new hires',
        ],
        'integrations': 'LinkedIn, Instagram, Twitter/X, Facebook, TikTok, Pinterest',
        'faqs': [
            ('Which platforms do you support?', 'LinkedIn, Instagram, Twitter/X, Facebook Pages/Groups, TikTok, Pinterest, YouTube community posts, and Google Business Profile.'),
            ('Can it write captions in my brand voice?', 'Yes — connect the AI Content Generator to the Scheduler. Captions are generated in your brand tone automatically.'),
            ('How does optimal timing work?', 'The AI analyses your historical post performance and audience activity patterns per platform to identify the highest-engagement windows.'),
            ('Can I manage multiple accounts?', 'Yes. Unlimited accounts in the Agency plan. Switch between them from a single dashboard.'),
            ('What happens if a post fails?', 'The system retries automatically and sends you a notification. Failed posts are queued for manual review.'),
        ],
    },
    {
        'name': 'AI Photo Editor', 'slug': 'ai-photo-editor', 'category': 'Content AI',
        'tagline': 'Professional visuals. Zero design skills required.',
        'description': 'Transform raw photos into publish-ready visuals in 90 seconds. Remove backgrounds, apply brand styles, batch-process product images, and generate on-brand thumbnails — without Photoshop or a designer.',
        'stats': ['90 sec Average edit time', '40+ Styles and filters', '5× Cheaper than design agencies'],
        'features': ['AI Background Removal', 'Style Transfer', 'Batch Processing', 'Brand Kit Application'],
        'how_it_works': ['Upload your image', 'Select style or edit type', 'Download ready-to-publish asset'],
        'ai_role': 'Detects subjects, removes/replaces backgrounds, applies brand colours and typography, generates variants, and exports in platform-correct dimensions.',
        'human_role': 'Selects the style direction, reviews final assets, and applies creative judgement on brand fit.',
        'use_cases': [
            'E-Commerce — consistent product photos across 5,000 SKUs',
            'Content Creator — branded thumbnails in 90 seconds',
            'Marketing Team — social graphics without a designer',
        ],
        'integrations': 'Canva, Adobe CC, Shopify, Instagram, WordPress',
        'faqs': [
            ('What image formats are supported?', 'JPEG, PNG, WebP, HEIC, and RAW formats. Output in PNG, JPEG, or WebP at any resolution.'),
            ('Can it handle batch processing?', 'Yes. Upload up to 500 images at once. The AI processes them all with consistent style and branding.'),
            ('How accurate is the background removal?', '99.2% accuracy on product photos. Hair and complex edges are handled with the same precision as professional editors.'),
            ('Can I save my brand kit?', 'Yes. Save brand colours, fonts, logo placement rules, and padding settings. Applied automatically to every image.'),
            ('Does it work for video thumbnails?', 'Yes. Optimised presets for YouTube (1280×720), LinkedIn, Instagram, and TikTok thumbnails.'),
        ],
    },
    {
        'name': 'Workflow Generator', 'slug': 'workflow-generator', 'category': 'Content AI',
        'tagline': 'Automate anything. No code needed.',
        'description': 'Describe what you want to automate in plain English. The AI builds the workflow — connecting your apps, setting triggers, and handling errors — then deploys it live in minutes.',
        'stats': ['85% Reduction in manual tasks', '200+ App integrations', '15 min To first automation'],
        'features': ['Visual Drag-and-Drop Builder', 'Pre-built Templates', 'Multi-step Triggers', 'Error Handling'],
        'how_it_works': ['Describe or choose a trigger', 'Build your workflow visually', 'Activate and monitor'],
        'ai_role': 'Interprets plain-English automation descriptions, maps them to app actions, handles conditional logic, and generates error-recovery paths.',
        'human_role': 'Defines business logic, approves automation scope, monitors exception logs.',
        'use_cases': [
            'Agency — auto-deliver reports to clients when campaigns end',
            'E-Commerce — trigger fulfilment workflow on every new order',
            'HR Team — onboarding workflow runs automatically on first login',
        ],
        'integrations': 'Zapier, Make, Slack, Google Workspace, HubSpot',
        'faqs': [
            ('Do I need coding skills?', 'None. Describe what you want in plain English and the AI builds it. Advanced users can edit the underlying JSON directly.'),
            ('How many steps can a workflow have?', 'Unlimited steps. Complex multi-branch workflows with parallel paths are fully supported.'),
            ('What triggers are available?', 'Webhooks, schedules (cron), form submissions, email events, CRM updates, and API calls from any connected app.'),
            ('Can workflows run on a schedule?', 'Yes. Set any cron schedule or trigger manually via API. Run every minute or once a month — your choice.'),
            ('What happens if a step fails?', 'The workflow retries failed steps up to 3 times with exponential backoff. After that, you receive a failure notification with full logs.'),
        ],
    },
    # SALES
    {
        'name': 'Lead Scraper', 'slug': 'lead-scraper', 'category': 'Sales AI',
        'tagline': 'Find your next 1,000 customers — automatically.',
        'description': 'Define your ideal customer profile and let the AI scrape, verify, score, and deliver a ready-to-contact lead list in hours — not weeks. Pull from LinkedIn, company databases, job boards, and 50+ sources simultaneously.',
        'stats': ['50K+ Leads per day capacity', '97% Data accuracy', '10× Faster than manual research'],
        'features': ['Multi-source Scraping', 'AI Qualification Scoring', 'CRM Export', 'Real-time Verification'],
        'how_it_works': ['Define your ideal customer profile', 'AI scrapes and scores leads', 'Export verified list to CRM'],
        'ai_role': 'Searches 50+ data sources, extracts contact details, verifies emails and phones in real time, scores each lead against your ICP, and removes duplicates.',
        'human_role': 'Defines ICP criteria, reviews top-scored leads, approves export, and runs outreach.',
        'use_cases': [
            'SaaS outbound — 500 qualified leads delivered daily to HubSpot',
            'Recruiter — candidate sourcing at 10× speed',
            'Agency — client prospecting across 5 verticals simultaneously',
        ],
        'integrations': 'Salesforce, HubSpot, LinkedIn, Apollo, Pipedrive',
        'faqs': [
            ('Where does the lead data come from?', 'LinkedIn, company websites, job boards, news mentions, business directories, and proprietary databases — 50+ sources in total.'),
            ('How is data accuracy maintained?', 'Every email is verified in real time against mail servers. Phone numbers are validated against carrier databases. 97% accuracy guaranteed.'),
            ('Is this GDPR compliant?', 'Yes. We only collect publicly available business contact data and provide full data processing documentation for GDPR compliance.'),
            ('Can I set ICP filters?', 'Yes. Filter by industry, company size, job title, geography, tech stack, funding stage, growth signals, and 40+ more attributes.'),
            ('How does AI scoring work?', 'The model scores each lead 0–100 based on how closely they match your ICP definition. You set the minimum threshold for export.'),
        ],
    },
    {
        'name': 'WhatsApp Automation', 'slug': 'whatsapp-automation', 'category': 'Sales AI',
        'tagline': 'Your sales team. Available 24/7.',
        'description': 'Build intelligent WhatsApp conversation flows that qualify leads, answer questions, book meetings, and hand off hot prospects to your sales team — all without a human touching a keyboard.',
        'stats': ['98% Message open rate', '3× Conversion vs email', '24/7 Automated engagement'],
        'features': ['AI Conversation Flows', 'Broadcast Campaigns', 'Lead Qualification', 'CRM Handoff'],
        'how_it_works': ['Build your conversation flow', 'Set triggers and conditions', 'AI engages and qualifies 24/7'],
        'ai_role': 'Manages the full conversation — sends messages, interprets replies, asks qualification questions, books calendar slots, and triggers CRM updates.',
        'human_role': 'Designs conversation strategy, sets qualification criteria, handles hot leads after AI handoff.',
        'use_cases': [
            'Real estate agency — qualifies buyers before agent contact',
            'E-Commerce — abandoned cart recovery with 98% open rate',
            'B2B SaaS — books demos automatically from inbound leads',
        ],
        'integrations': 'WhatsApp Business API, HubSpot, Salesforce, Zapier, Twilio',
        'faqs': [
            ('Do I need the WhatsApp Business API?', 'Yes. We help you get approved and set up during onboarding — typically takes 24–48 hours.'),
            ('Can I send broadcast messages?', 'Yes. Send to opted-in contacts only (WhatsApp policy). Personalise with name, company, and custom fields.'),
            ('How does the AI handle unexpected replies?', 'The AI uses intent detection to route unexpected messages to the closest relevant flow. Truly unrecognised inputs escalate to a human.'),
            ('Can it book meetings in my calendar?', 'Yes. Connects with Calendly, Google Calendar, and Cal.com. The AI checks availability and confirms bookings in conversation.'),
            ('Is this compliant with WhatsApp policies?', 'Yes. We only message opted-in contacts, respect opt-out requests instantly, and follow all WhatsApp Business policy guidelines.'),
        ],
    },
    {
        'name': 'Cold Calling Bot', 'slug': 'cold-calling-bot', 'category': 'Sales AI',
        'tagline': '1,000 calls a day. Zero burnout.',
        'description': 'Deploy an AI voice agent that cold calls your prospect list, delivers a natural conversation, handles objections, and books qualified meetings — transferring hot calls to your team live.',
        'stats': ['1,000+ Calls per day', '40% Connect rate', '8× More qualified meetings'],
        'features': ['Natural AI Voice', 'Objection Handling', 'Live Transfer', 'Call Analytics'],
        'how_it_works': ['Upload your prospect list', 'AI calls and qualifies', 'Hot leads transferred live to your team'],
        'ai_role': 'Dials prospects, delivers opening, adapts to conversation flow, handles standard objections, qualifies interest, and initiates live transfer or books a callback.',
        'human_role': 'Records opening script, sets qualification criteria, takes live-transferred calls, reviews call recordings for coaching.',
        'use_cases': [
            'Insurance broker — 8× more qualified meetings per week',
            'SaaS company — 1,000 outbound calls daily with 2-person team',
            'Recruitment firm — candidate screening at scale',
        ],
        'integrations': 'Salesforce, HubSpot, Twilio, Zoom, Calendly',
        'faqs': [
            ('Does it sound like a real person?', 'Our voice AI uses the latest neural TTS. In blind tests, 78% of recipients cannot distinguish it from a human caller.'),
            ('How does it handle objections?', 'You pre-define objection responses during setup. The AI detects objection patterns and delivers the appropriate response naturally.'),
            ('Can it transfer to a human?', 'Yes. When a prospect is qualified and interested, the AI says "let me connect you with a specialist now" and bridges the call live.'),
            ('What is the connect rate?', 'Average 40% connect rate from our customer base. Varies by industry, time of day, and list quality.'),
            ('Is cold calling legal?', 'Compliant calling requires following local TCPA/DNC regulations. Our system includes DNC list scrubbing and compliance checks before every dial.'),
        ],
    },
    {
        'name': 'CRM Sync', 'slug': 'crm-sync', 'category': 'Sales AI',
        'tagline': 'Your CRM, always up to date.',
        'description': 'Eliminate manual data entry forever. CRM Sync automatically keeps your contacts, deals, activities, and notes in perfect sync across all your tools — bidirectionally, in real time.',
        'stats': ['100% Data accuracy', 'Real-time sync', '3hrs Saved per rep per day'],
        'features': ['Bi-directional Sync', 'AI Data Enrichment', 'Duplicate Detection', 'Activity Logging'],
        'how_it_works': ['Connect your CRM and tools', 'Map your data fields', 'AI syncs continuously in real time'],
        'ai_role': 'Monitors all connected systems for changes, resolves conflicts, enriches records with public data, detects and merges duplicates, and logs all activities automatically.',
        'human_role': 'Configures sync rules, reviews conflict resolution logic, handles edge cases flagged by the system.',
        'use_cases': [
            'Sales team — zero data entry, full activity logging from email and calls',
            'RevOps — single source of truth across Salesforce and HubSpot',
            'Agency — client data synced across 6 different tools automatically',
        ],
        'integrations': 'Salesforce, HubSpot, Pipedrive, Notion, Airtable',
        'faqs': [
            ('Which CRMs are supported?', 'Salesforce, HubSpot, Pipedrive, Zoho, Monday.com, Notion, and Airtable. Custom CRMs supported via our REST API.'),
            ('How does conflict resolution work?', 'By default, the most recently updated record wins. You can configure custom rules per field — e.g., always trust Salesforce for deal stage.'),
            ('Is historical data migrated?', 'Yes. On setup, we sync your full historical record going back to account creation, with de-duplication included.'),
            ('What about custom fields?', 'Full custom field mapping supported. Map any field in System A to any field in System B, with value transformation rules.'),
            ('How quickly does sync happen?', 'Changes are synced within 60 seconds of being made. Critical fields like deal stage sync in under 10 seconds.'),
        ],
    },
    # RESEARCH
    {
        'name': 'Competitor Analyzer', 'slug': 'competitor-analyzer', 'category': 'Research AI',
        'tagline': 'Know everything. Act faster.',
        'description': 'Monitor every move your competitors make — pricing changes, new features, content strategy, hiring signals, ad campaigns, and customer sentiment — and receive actionable intelligence in a weekly digest.',
        'stats': ['500+ Data points per competitor', 'Updated daily', '6hrs Saved per week'],
        'features': ['Real-time Monitoring', 'Pricing Tracking', 'Content Gap Analysis', 'SWOT Generation'],
        'how_it_works': ['Add your competitors by URL or name', 'AI monitors all channels daily', 'Receive weekly intelligence reports'],
        'ai_role': 'Crawls competitor websites, social channels, job boards, review platforms, and ad libraries daily. Extracts signals, identifies trends, and generates SWOT analysis automatically.',
        'human_role': 'Reviews weekly digests, decides strategic responses, and shares intelligence with relevant teams.',
        'use_cases': [
            'SaaS company — caught competitor price drop 3hrs after it happened',
            'E-Commerce — identified untapped content gap that drove 40% SEO traffic increase',
            'Agency — delivers competitive intelligence reports to 15 clients weekly',
        ],
        'integrations': 'SEMrush, Ahrefs, Google Analytics, Slack, Notion',
        'faqs': [
            ('How many competitors can I track?', 'Up to 10 on Standard, unlimited on Pro. We recommend tracking 5–7 direct competitors for actionable insights.'),
            ('What data sources are monitored?', 'Website changes, blog posts, social media, job postings, G2/Trustpilot reviews, Google Ads, Meta Ads, and pricing pages.'),
            ('How quickly are changes detected?', 'Website and pricing changes detected within 24 hours. Social posts within 2 hours. Ad changes within 48 hours.'),
            ('Can I get alerts for specific changes?', 'Yes. Set custom alerts for pricing changes, new feature announcements, or any keyword appearing on their site.'),
            ('Is the SWOT analysis accurate?', 'The AI-generated SWOT is based on real data signals. We recommend reviewing and annotating it with internal context before sharing with stakeholders.'),
        ],
    },
    {
        'name': 'Pricing Intelligence', 'slug': 'pricing-intelligence', 'category': 'Research AI',
        'tagline': 'Always price to win.',
        'description': 'Get real-time visibility into competitor pricing across every channel, with AI-powered recommendations for your optimal price point — balancing margin, conversion rate, and market position.',
        'stats': ['Real-time price tracking', '23% Average margin improvement', '100+ Markets covered'],
        'features': ['Dynamic Price Recommendations', 'Competitor Price Alerts', 'Elasticity Modelling', 'A/B Test Tracking'],
        'how_it_works': ['Connect your product catalogue', 'AI tracks market prices in real time', 'Receive daily price recommendations'],
        'ai_role': 'Tracks competitor prices across all channels, models demand elasticity, calculates optimal price points, runs A/B test analysis, and generates pricing action recommendations.',
        'human_role': 'Reviews recommendations, approves price changes, monitors margin impact, and sets pricing guardrails.',
        'use_cases': [
            'E-Commerce retailer — 23% margin improvement in 60 days',
            'SaaS company — optimised 3 pricing tiers, reduced churn 18%',
            'Marketplace seller — reprices 8,000 SKUs daily automatically',
        ],
        'integrations': 'Shopify, WooCommerce, Stripe, Google Shopping, Amazon',
        'faqs': [
            ('How often are prices tracked?', 'Competitor prices are checked every 4 hours. For high-velocity markets, real-time tracking (every 15 minutes) is available on Enterprise.'),
            ('Does it automatically change my prices?', 'Only if you enable autopilot mode with guardrails (min/max price limits). By default, all recommendations require human approval.'),
            ('What is elasticity modelling?', 'The AI estimates how demand changes at different price points using your historical sales data, so recommendations optimise revenue, not just margin.'),
            ('Can it handle bundles and tiered pricing?', 'Yes. Bundle pricing, volume discounts, and tiered plans are all modelled. Each configuration gets its own optimisation.'),
            ('Which markets does it cover?', '100+ countries and all major e-commerce platforms. Pricing is tracked in local currency with FX normalisation for comparison.'),
        ],
    },
    {
        'name': 'Market Strategy Generator', 'slug': 'market-strategy-generator', 'category': 'Research AI',
        'tagline': 'Your GTM strategy. In 48 hours.',
        'description': 'Input your product, target market, and competitive context. The AI researches 200+ data sources, analyses the opportunity, and delivers a complete go-to-market strategy — positioning, personas, channel mix, pricing, and messaging.',
        'stats': ['48h Full strategy delivery', '200+ Data sources', 'Used by Fortune 500 teams'],
        'features': ['Market Sizing', 'Persona Generation', 'Channel Strategy', 'Competitive Positioning'],
        'how_it_works': ['Input your product and target market', 'AI researches and analyses 200+ sources', 'Receive complete GTM strategy as a PDF'],
        'ai_role': 'Performs market sizing (TAM/SAM/SOM), generates 3–5 detailed buyer personas, recommends optimal channel mix, writes positioning statement and messaging framework.',
        'human_role': 'Provides product context and strategic constraints, validates persona accuracy with internal knowledge, selects and executes the recommended strategy.',
        'use_cases': [
            'Startup — full GTM before seed pitch, secured funding in 3 weeks',
            'Enterprise — new market entry strategy for 3 countries in 48 hours',
            'Agency — delivers competitive strategy reports to clients at 10× speed',
        ],
        'integrations': 'Google Analytics, SEMrush, LinkedIn, HubSpot, Notion',
        'faqs': [
            ('What is included in the strategy output?', 'TAM/SAM/SOM analysis, 3–5 buyer personas, competitive landscape, positioning statement, channel recommendations, messaging framework, and a 90-day execution plan.'),
            ('How is the market data sourced?', 'We aggregate from industry reports, government databases, LinkedIn data, web scraping, SEMrush, and proprietary sources — 200+ in total.'),
            ('Can I customise the strategy template?', 'Yes. White-label the output with your branding. Add or remove sections. Export as PDF, Word, or slide deck.'),
            ('How accurate is the market sizing?', 'TAM/SAM figures are drawn from authoritative sources and cross-validated. We include data source citations so you can verify every number.'),
            ('What if my market is very niche?', 'Our AI has coverage across 500+ industries and verticals. For highly specialised markets, we supplement automated research with curated data.'),
        ],
    },
    # BUSINESS
    {
        'name': 'CRM System', 'slug': 'crm-system', 'category': 'Business AI',
        'tagline': 'The CRM that sells for you.',
        'description': 'A fully AI-powered CRM that scores deals, writes follow-up emails, forecasts revenue, and tells your team exactly what to do next — so no deal falls through the cracks.',
        'stats': ['360° Customer view', '45% Faster deal cycles', 'AI-powered next-action suggestions'],
        'features': ['AI Deal Scoring', 'Pipeline Automation', 'Email Sequences', 'Revenue Forecasting'],
        'how_it_works': ['Import your contacts and deals', 'AI scores, segments, and prioritises', 'Automated sequences run while you focus on closing'],
        'ai_role': 'Scores every deal 0–100, suggests next best action, writes personalised follow-up emails, forecasts monthly revenue, and identifies at-risk deals before they churn.',
        'human_role': 'Reviews AI suggestions, takes priority calls, approves email sequences, and makes final deal decisions.',
        'use_cases': [
            'Sales team of 5 — managing pipeline of 200 active deals with zero dropped follow-ups',
            'Agency — automated client reporting and renewal reminders',
            'SaaS company — 45% faster deal cycles with AI next-action prompts',
        ],
        'integrations': 'Gmail, Outlook, Slack, Stripe, Zapier',
        'faqs': [
            ('How does AI deal scoring work?', 'The model analyses engagement signals (email opens, meeting attendance, response times), deal characteristics, and historical patterns to score likelihood of closing.'),
            ('Can it write emails in my voice?', 'Yes. Train it on 10+ of your best-performing emails. All generated follow-ups match your tone, length, and style.'),
            ('Does it integrate with my email?', 'Full Gmail and Outlook integration. Emails sent from the CRM appear in your sent folder. Replies are automatically logged.'),
            ('How accurate is revenue forecasting?', 'Our models achieve 85–92% forecast accuracy within a 30-day window, based on historical data from 500+ customer accounts.'),
            ('Can I migrate from my current CRM?', 'Yes. We offer guided migration from Salesforce, HubSpot, Pipedrive, and any CRM with CSV export. Historical data and attachments are preserved.'),
        ],
    },
    {
        'name': 'ERP Modules', 'slug': 'erp-modules', 'category': 'Business AI',
        'tagline': 'Run your business. On autopilot.',
        'description': 'A modular AI-powered ERP covering finance, inventory, HR, payroll, and project management — all in one connected system. Real-time dashboards replace spreadsheets. Automated workflows replace manual processes.',
        'stats': ['60% Reduction in admin overhead', 'All departments connected', 'Real-time dashboards'],
        'features': ['Finance Module', 'Inventory Management', 'HR & Payroll', 'Project Management'],
        'how_it_works': ['Connect your business units and data sources', 'AI centralises and normalises data', 'Single dashboard for all operational decisions'],
        'ai_role': 'Automates invoice processing, payroll calculations, inventory reorder triggers, project timeline updates, and generates management dashboards in real time.',
        'human_role': 'Approves large transactions, makes strategic staffing decisions, reviews exception reports, and sets business rules.',
        'use_cases': [
            'Manufacturer — eliminated 3 separate tools, saved £4,200/month in software costs',
            'Services firm — payroll now runs automatically on the 25th of each month',
            'Retailer — inventory reorder triggers prevent stockouts across 12 locations',
        ],
        'integrations': 'QuickBooks, Xero, Stripe, Slack, Google Workspace',
        'faqs': [
            ('Do I need all modules or can I start with one?', 'Start with any single module. Add more as you grow. Pricing is per module per month — no all-or-nothing commitment.'),
            ('How long does implementation take?', 'A single module (e.g. Finance) goes live in 3–5 days. Full multi-module implementation typically takes 2–3 weeks with our onboarding team.'),
            ('Can it handle multi-currency?', 'Yes. Full multi-currency support with real-time FX rates. Reports can be normalised to any base currency.'),
            ('Does it replace my accountant?', 'No — it handles the repetitive processing work so your accountant focuses on strategy and compliance, not data entry.'),
            ('Is payroll compliant with local law?', 'Payroll is compliant for UK, US, EU, and India. Additional regions are added regularly. Check our compliance page for current coverage.'),
        ],
    },
    {
        'name': 'LMS System', 'slug': 'lms-system', 'category': 'Business AI',
        'tagline': 'Train your team. Scale your knowledge.',
        'description': 'An AI-powered learning management system that turns your SOPs, guides, and expertise into structured training courses in minutes — with automated assessments, personalised learning paths, and completion tracking.',
        'stats': ['10× Faster course creation', '94% Completion rates', 'AI-personalised learning paths'],
        'features': ['AI Course Builder', 'Progress Tracking', 'Certification', 'Team Analytics'],
        'how_it_works': ['Upload your training content or SOPs', 'AI structures it into courses with quizzes', 'Assign to team — AI personalises the path'],
        'ai_role': 'Converts raw documents into structured lessons, generates quiz questions, personalises learning order based on role and prior knowledge, and identifies knowledge gaps in your team.',
        'human_role': 'Provides source material, reviews course structure, approves certification criteria, and acts on team analytics.',
        'use_cases': [
            'SaaS company — onboarded 50 new hires in one week with automated training',
            'Franchise — consistent training across 40 locations with no trainer travel',
            'Consulting firm — knowledge base turned into billable training product',
        ],
        'integrations': 'Zoom, Google Workspace, Slack, HubSpot, Stripe',
        'faqs': [
            ('What content formats can I upload?', 'PDFs, Word docs, PowerPoint, video files (MP4), audio, and URLs. The AI extracts and structures content from all formats.'),
            ('Can learners complete on mobile?', 'Yes. The learner portal is fully mobile-responsive. Courses can be downloaded for offline access on iOS and Android.'),
            ('How does AI personalisation work?', 'The AI assesses each learner\'s existing knowledge, role, and learning pace — then reorders and adjusts content to close gaps most efficiently.'),
            ('Can I sell courses externally?', 'Yes. Enable the public storefront, set pricing, and sell via Stripe. Great for turning internal training into a revenue stream.'),
            ('Are certificates verifiable?', 'Yes. Each certificate has a unique verification URL. Certificates can be shared on LinkedIn directly from the learner\'s portal.'),
        ],
    },
    {
        'name': 'Website Builder', 'slug': 'website-builder', 'category': 'Business AI',
        'tagline': 'Beautiful websites. No developers needed.',
        'description': 'Build a professional, SEO-optimised website in 30 minutes using AI-generated copy, 100+ conversion-optimised templates, and a visual drag-and-drop editor — then publish to a custom domain with one click.',
        'stats': ['Launch in 30 minutes', '100+ Templates', 'SEO-optimised by default'],
        'features': ['Drag-and-Drop Builder', 'AI Copy Generation', 'Mobile-first Design', 'CMS Built-in'],
        'how_it_works': ['Choose a template that fits your business', 'AI generates all copy from your brief', 'Publish to your custom domain in one click'],
        'ai_role': 'Generates all website copy from a brief, selects and places images, optimises heading structure and meta tags for SEO, and suggests layout improvements based on conversion data.',
        'human_role': 'Selects template and brand direction, reviews and edits copy, uploads logo and brand assets, and approves before launch.',
        'use_cases': [
            'Freelancer — professional website live in 30 minutes before a client pitch',
            'Local business — replaces £3,000/year web agency with a £49/month subscription',
            'SaaS startup — landing page live and collecting emails before MVP is built',
        ],
        'integrations': 'Custom domain, Google Analytics, HubSpot, Stripe, Mailchimp',
        'faqs': [
            ('Can I use my own domain?', 'Yes. Connect any custom domain purchased from any registrar. We provide the DNS settings and SSL certificate automatically.'),
            ('Is e-commerce included?', 'Basic product pages and Stripe payment links are included. For full e-commerce (inventory, multi-SKU, cart), use the E-Commerce Sites tool.'),
            ('How good is the AI-generated copy?', 'The AI writes based on your brief and is trained on high-converting website copy. 80% of customers publish with minimal editing. 20% use it as a first draft.'),
            ('Can I edit the code?', 'Yes. Access the HTML/CSS of any element directly. Or use the visual editor exclusively — your choice.'),
            ('What happens to my site if I cancel?', 'You can export your site as a static HTML/CSS package at any time. Your content is always yours.'),
        ],
    },
    {
        'name': 'E-Commerce Sites', 'slug': 'ecommerce-sites', 'category': 'Business AI',
        'tagline': 'Your store. Selling while you sleep.',
        'description': 'Launch a fully-featured AI-powered e-commerce store in 30 minutes. AI writes your product listings, recommends upsells, recovers abandoned carts, and syncs inventory across channels — so your store sells itself.',
        'stats': ['30 min Setup to live store', 'AI product recommendations', 'Multi-currency support'],
        'features': ['AI Product Listings', 'Abandoned Cart Recovery', 'Inventory Sync', 'Multi-channel Selling'],
        'how_it_works': ['Add your products and connect payment', 'AI writes listings and sets up recovery flows', 'Go live — AI manages recommendations and sync'],
        'ai_role': 'Writes SEO product titles and descriptions, generates upsell/cross-sell recommendations, sends abandoned cart emails, syncs inventory across Shopify/Amazon/Instagram, and analyses sales data.',
        'human_role': 'Sources and photographs products, sets pricing strategy, manages supplier relationships, and makes promotional decisions.',
        'use_cases': [
            'First-time seller — store live with 50 products in under 2 hours',
            'Established brand — added 22% revenue from AI-powered upsell recommendations',
            'Multi-channel retailer — inventory synced across 4 channels with zero overselling',
        ],
        'integrations': 'Stripe, PayPal, Shopify, Amazon, Instagram Shopping',
        'faqs': [
            ('What payment methods are supported?', 'Stripe (cards, Apple Pay, Google Pay, BNPL), PayPal, and bank transfer. Crypto payments available on request.'),
            ('How does abandoned cart recovery work?', 'The AI sends a sequence of 2–3 personalised emails to customers who added items but didn\'t complete checkout. Average recovery rate: 12–18%.'),
            ('Can I sell digital products?', 'Yes. Digital downloads, subscription access, and course access are all supported with automatic delivery on purchase.'),
            ('How does multi-channel sync work?', 'Connect Shopify, Amazon, Instagram, and your Agentix store. Inventory decrements across all channels simultaneously when a sale is made.'),
            ('Is there a transaction fee?', 'No Agentix transaction fee. You pay only the payment processor fee (Stripe: 1.4% + 20p for European cards). We never take a cut of your sales.'),
        ],
    },
]

for tool in tools_data:
    h2(f"{tool['name']}  —  /tools/{tool['slug']}")
    label_val('Category', tool['category'])
    label_val('Tagline', tool['tagline'])
    label_val('Description', tool['description'])

    p = doc.add_paragraph()
    p.add_run('Stats: ').bold = True
    p.add_run(' | '.join(tool['stats']))

    p = doc.add_paragraph()
    p.add_run('Features: ').bold = True
    p.add_run(', '.join(tool['features']))

    p = doc.add_paragraph()
    p.add_run('How It Works: ').bold = True
    p.add_run(' → '.join(tool['how_it_works']))

    label_val('AI Role', tool['ai_role'])
    label_val('Human Role', tool['human_role'])

    p = doc.add_paragraph()
    p.add_run('Use Cases: ').bold = True
    for uc in tool['use_cases']:
        doc.add_paragraph(uc, style='List Bullet')

    label_val('Integrations', tool['integrations'])

    doc.add_paragraph('FAQs:').runs[0].bold = True
    for q, a in tool['faqs']:
        p = doc.add_paragraph()
        p.add_run(f'Q: {q}').bold = True
        doc.add_paragraph(f'A: {a}')

    doc.add_paragraph('')

doc.add_page_break()

# ── ABOUT ──────────────────────────────────────────────────────────
h1('ABOUT PAGE  (/about)')

h2('Hero')
label_val('Headline', 'Built for the AI-first business era')
label_val('Subheadline', 'We believe every company — regardless of size — deserves access to enterprise-grade AI tools. Agentix makes that possible.')

h2('Mission')
body('We started Agentix because we kept seeing the same problem: enterprise companies paying millions for custom AI, while smaller businesses were left behind. We built a marketplace of best-in-class AI tools that any business can deploy in minutes — without consultants, without code, without a data science team.')
body('Our mission is to democratise AI automation. Agentix gives every business the same tools that the Fortune 500 used to have exclusive access to.')

h2('Stats')
for stat in ['16 AI Tools', '4 Categories', '500+ Businesses', '99.9% Uptime']:
    bullet(stat)

h2('Values')
for val in ['AI for Everyone — enterprise power, startup accessibility', 'Speed to Value — live in 30 minutes, not 3 months', 'Human + AI — augmentation, not replacement', 'Relentless Improvement — every tool gets better every week']:
    bullet(val)

h2('Team Philosophy')
body('Agentix is a globally distributed, remote-first team. We are AI-native, customer-obsessed, and relentlessly focused on making our tools produce measurable results for every customer.')

doc.add_page_break()

# ── CONTACT ────────────────────────────────────────────────────────
h1('CONTACT PAGE  (/contact)')

label_val('Email', 'hello@ai-agentix.com')
label_val('Press', 'press@ai-agentix.com')
label_val('Book a Demo', 'Via contact form on /contact')

h2('Form Fields')
for field in ['Full Name', 'Work Email', 'Company', 'Phone (optional)', 'Message / What are you looking to automate?']:
    bullet(field)

h2('Contact Page FAQs')
for q, a in [
    ('How quickly will you respond?', 'Within 24 hours on business days. Demo requests are prioritised and typically responded to within 4 hours.'),
    ('Is there a free trial?', 'Yes. Every tool has a 14-day free trial. No credit card required.'),
    ('Do you offer custom enterprise plans?', 'Yes. Contact us for volume pricing, custom integrations, and dedicated support SLAs.'),
]:
    p = doc.add_paragraph()
    p.add_run(f'Q: {q}').bold = True
    doc.add_paragraph(f'A: {a}')

doc.add_page_break()

# ── USE CASES ──────────────────────────────────────────────────────
h1('USE CASES PAGE  (/use-cases)')

use_cases = [
    ('Marketing Agency — Content AI', '12× content output, same team size', 'Used AI Content Generator + Social Media Scheduler to produce content for 25 clients simultaneously.'),
    ('SaaS Sales Team — Sales AI', '3× more qualified leads', 'Deployed Lead Scraper + Cold Calling Bot for 1,000 automated outreach calls per day.'),
    ('E-Commerce Brand — Research AI', '18% margin improvement', 'Used Pricing Intelligence + Competitor Analyzer to reprice 8,000 SKUs daily in real time.'),
    ('Retail Chain — Business AI', '60% operational cost reduction', 'Unified CRM, ERP, and inventory across 12 locations under one Agentix dashboard.'),
    ('Personal Brand — Content AI', '4.2× follower growth in 90 days', 'Social Media Scheduler + AI Content Generator in 30 minutes per week.'),
    ('B2B Consultancy — Sales AI', '2× deal close rate', 'WhatsApp Automation + CRM Sync eliminated all manual follow-up.'),
    ('Product Launch — Research AI', 'Full GTM strategy in 48 hours', 'Market Strategy Generator delivered positioning, personas, and channel mix before first sales call.'),
    ('Online Store — Business AI', 'Live e-commerce in 30 minutes', 'E-Commerce Sites + Website Builder — from idea to payment-enabled store.'),
]

for title, result, desc in use_cases:
    h3(title)
    label_val('Result', result)
    body(desc)

doc.add_page_break()

# ── RESOURCES ──────────────────────────────────────────────────────
h1('RESOURCES PAGE  (/resources)')

resources = [
    ('Guide', 'Getting Started with Agentix Content AI', 'Step-by-step setup for AI Content Generator, Social Media Scheduler, and Workflow Generator.', '24 pages'),
    ('Guide', 'The Complete AI Sales Playbook', 'How to use Lead Scraper, WhatsApp Automation, and Cold Calling Bot together for a fully automated pipeline.', '36 pages'),
    ('Whitepaper', 'ROI of AI Tools for SMBs in 2026', 'Data from 500+ Agentix customers on time saved, revenue generated, and costs reduced.', '28 pages'),
    ('Webinar', 'Live Demo: Agentix Sales AI in Action', 'Full Sales AI suite demo — lead generation to CRM sync — with live Q&A.', '60 min'),
    ('Guide', 'Market Research AI: Zero to Insights in 48 Hours', 'Use Competitor Analyzer + Pricing Intelligence for complete competitive intelligence.', '20 pages'),
    ('Ebook', 'Agentix Business AI: Complete Onboarding Guide', '30-step checklist for CRM, ERP, LMS, Website Builder, and E-Commerce tools.', '44 pages'),
]

for rtype, title, desc, length in resources:
    h3(f'[{rtype}] {title}')
    label_val('Description', desc)
    label_val('Length', length)

doc.add_page_break()

# ── PRESS ──────────────────────────────────────────────────────────
h1('PRESS PAGE  (/press)')

press = [
    ('January 2026', 'Agentix Launches 16-Tool AI Marketplace for SMBs', 'Agentix today announced the launch of its comprehensive AI tools platform, bringing enterprise-grade automation to businesses of all sizes across content, sales, research, and operations.'),
    ('February 2026', 'Agentix Content AI Helps Marketing Teams 10× Output', 'Early adopters of the Agentix Content AI suite report producing 10 times more content with the same team, using the AI Content Generator and Social Media Scheduler.'),
    ('March 2026', 'Agentix Raises Strategic Investment to Accelerate AI Tool Development', 'Agentix secured strategic investment to expand its 16-tool marketplace, with upcoming additions across the Sales AI and Business AI categories.'),
]

for date, headline, excerpt in press:
    h3(headline)
    label_val('Date', date)
    body(excerpt)

h2('Media Coverage')
for pub in ['TechCrunch', 'Forbes', 'VentureBeat', 'Wired']:
    bullet(pub)

label_val('Press Contact', 'press@ai-agentix.com')

# ── SAVE ───────────────────────────────────────────────────────────
output_path = r'd:\AI AGENTIX\Agentix-Website-Content.docx'
doc.save(output_path)
print(f'Saved: {output_path}')

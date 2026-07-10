# -*- coding: utf-8 -*-
"""Generate the AI Agentix Master Image Asset Requirement Document (.docx)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE = "F26522"; NAVY = "0B1F3A"; BLUE = "2F80ED"; LIGHT = "F1F1F3"; DARK = "12141A"; WHITE = "FFFFFF"

COLORS = ("Primary Orange #F26522 with secondary blue accents (#2F80ED / deep navy #0B1F3A) "
          "over a dark, cinematic base. Must communicate trust, innovation, healthcare/tech and "
          "professionalism. No off-brand colours that break site consistency.")
LIB_STYLE = ("Photorealistic 'AI-in-business' scene blended with subtle holographic UI / dashboard "
             "elements. Premium, modern, enterprise. Must match the approved Home Page image style.")
HERO_STYLE = ("Cinematic, wide, premium hero shot. Photorealistic environment with subtle holographic "
              "AI/UI accents and an orange rim-light. Subject/action kept centre-left (headline sits on "
              "the left); right side breathes. Matches approved Home Page style.")

# ---------- low-level helpers ----------
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def set_text(cell, text, bold=False, color=None, size=9, align=None, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'TOC \\o "1-2" \\h \\z \\u')
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = "Right-click and choose 'Update Field' to build the Table of Contents."
    r.append(t); fld.append(r); p._p.append(fld)

# ---------- image-entry renderer ----------
def image_entry(doc, e):
    h = doc.add_heading(f"{e['id']} — {e['name']}", level=3)
    rows = [
        ('Page / Section', e['where']),
        ('Purpose', e['purpose']),
        ('Image Content Description', e['content']),
        ('Visual Style', e.get('style', LIB_STYLE)),
        ('Theme Colours', COLORS),
        ('Aspect Ratio', e['ratio']),
        ('Recommended Resolution', e['res']),
        ('Orientation', e['orient']),
        ('File Format', e.get('fmt', 'PNG or JPG (delivered) — will be optimised to WebP')),
        ('Priority', e['priority']),
        ('Reuse / Used On', e['reuse']),
        ('Notes for Designer', e['notes']),
    ]
    t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'; t.alignment = 1
    for k, v in rows:
        c = t.add_row().cells
        set_text(c[0], k, bold=True, size=9, color=DARK); shade(c[0], LIGHT)
        col = ORANGE if k == 'Priority' else None
        set_text(c[1], v, size=9, color=col, bold=(k == 'Priority'))
        c[0].width = Inches(1.7); c[1].width = Inches(4.9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def section_title(doc, num, title, blurb=None):
    doc.add_page_break()
    doc.add_heading(f"{num}. {title}", level=1)
    if blurb:
        p = doc.add_paragraph(blurb); p.runs[0].font.size = Pt(10.5); p.runs[0].italic = True

# =========================================================================
# DATA
# =========================================================================
LIB_DEFAULT = dict(where="Reusable content image — appears in Solution capability panels and Industry "
                         "solution rows across the site.",
                   style=LIB_STYLE, ratio="4:3", res="1600 × 1200 px", orient="Landscape",
                   priority="High")

def lib(id_, name, content, reuse, notes, **kw):
    e = dict(LIB_DEFAULT); e.update(dict(id=id_, name=name, content=content, reuse=reuse, notes=notes)); e.update(kw)
    e['purpose'] = e.get('purpose', f"Reusable visual representing “{name.replace('AI ', '')}”. One image "
                                    f"serves multiple pages so the design set stays small and consistent.")
    return e

LIBRARY = [
 lib("LIB-01", "AI Lead Qualification & Scoring",
     "An AI interface scoring and ranking inbound leads — cards/avatars with intent scores, a qualification funnel and a highlighted 'hot lead'. A professional reviews it on a large screen.",
     "Sales Automation (Lead Qualification AI) · Real Estate (Lead Qualification Engine) · AI Voice & Chat (Lead Qualification Engine)",
     "Keep the funnel/score UI legible; subject centre-left."),
 lib("LIB-02", "AI Omnichannel Outreach",
     "AI sending personalised outreach across WhatsApp, email and phone simultaneously — floating chat/mail/call bubbles connecting to contact avatars, orchestrated from one console.",
     "Sales Automation (Multi-Channel Outreach) · Real Estate (Site Visit Automation) · HRMS (Interview scheduling)",
     "Show 3 channels (WhatsApp green, email, call) unified — but keep brand orange dominant on UI."),
 lib("LIB-03", "AI Chat & Customer Support",
     "A modern support scene where an AI chat assistant resolves customer queries in real time; a friendly agent oversees a wall of live conversations with resolution stats.",
     "AI Voice & Chat (NLU) · Retail (Support Automation) · Education (Student Support Chatbot) · Operations (Exception Management)",
     "Convey speed + 24/7. Chat UI in brand colours."),
 lib("LIB-04", "AI Voice Agent",
     "A human-like AI voice agent handling calls — a headset/soundwave motif, live call transcript panel, calendar booking confirmation, multilingual tags.",
     "AI Voice & Chat (Multi-Language / Voice) · Hospital Management (Patient Communication)",
     "Voice waveform + transcript should read as 'natural, human-like'."),
 lib("LIB-05", "AI Analytics & BI Dashboard",
     "A premium real-time analytics command centre — KPI cards, trend charts, anomaly alerts and forecasts on large holographic screens, reviewed by an executive.",
     "Revenue Analytics · Ops Analytics · Supply-Chain Analytics · Market Intelligence · Financial Reporting · Marketing Analytics & ROI · Data Intelligence",
     "Most-reused image (6+ placements). Charts must look credible, not decorative."),
 lib("LIB-06", "AI Document Processing & OCR",
     "AI reading and extracting data from invoices, forms and contracts — a document being scanned into structured fields with validation ticks; paperless, automated.",
     "Finance (Invoice Automation) · Operations (Document Intelligence) · Real Estate (Document Processing) · Healthcare (Medical Documentation) · Hospital (EMR/EHR)",
     "Show 'days → minutes' idea; extracted-field highlights."),
 lib("LIB-07", "AI Workflow Automation & Orchestration",
     "A no-code workflow/orchestration canvas connecting many business tools into one automated flow — nodes, arrows and status pills flowing across a screen; bottleneck cleared.",
     "Operations (Workflow Orchestration) · Operations (Integration Fabric) · End-to-End Automation",
     "Node-graph must look like a real automation builder."),
 lib("LIB-08", "AI CRM & Data Sync",
     "A unified CRM where AI auto-enriches contacts and syncs data across systems — a customer 360 profile with enrichment badges and bidirectional sync arrows.",
     "Sales (CRM Auto-Enrichment) · AI Voice & Chat (CRM Integration) · Retail / Operations data-sync",
     "Convey 'zero manual data entry'."),
 lib("LIB-09", "AI Content Generation Studio",
     "An AI content studio producing blogs, social posts, ad copy and images at scale — multiple content drafts generating on screen with brand-voice controls.",
     "Marketing (Content Engine) · Education (Faculty Productivity) · Education (Alumni Engagement)",
     "Show multi-format output + brand-voice tuning."),
 lib("LIB-10", "AI Marketing Campaigns & Nurturing",
     "Automated multi-step campaigns nurturing leads across email/SMS/WhatsApp — a journey builder with branching sequences, A/B tests and conversion lift.",
     "Marketing (Campaign Manager) · Marketing (Lead Nurturing) · Retail (Cart Recovery) · Real Estate (Post-Sale) · Logistics (Customer Communication) · Hospitality (Reputation)",
     "Journey/sequence UI; personalisation emphasis."),
 lib("LIB-11", "AI Scheduling & Booking",
     "AI autonomously booking appointments/meetings — a smart calendar auto-filling slots, confirmations sent, no-show reminders; a relieved team.",
     "Sales (Meeting Booking) · Hospital (Patient Scheduling) · Hospitality (Reservation) · HRMS (Interview Automation)",
     "Calendar + confirmation motif; 24/7."),
 lib("LIB-12", "AI Forecasting & Demand Planning",
     "ML demand-forecasting — predictive curves, seasonality bands and confidence intervals guiding stock/fleet decisions; a planner acting on the forecast.",
     "Supply Chain (Demand Forecasting) · Logistics (Demand Forecasting) · Real Estate (Market Intelligence)",
     "Forecast chart must look predictive (future band)."),
 lib("LIB-13", "AI Inventory & Warehouse Intelligence",
     "A smart warehouse with AI tracking stock in real time — shelves/bins with live counts, reorder alerts, pick-path optimisation overlays; robots/handhelds.",
     "Supply Chain (Inventory Intelligence) · Logistics (Warehouse Intelligence) · Retail (Inventory) · Healthcare (Pharmacy Inventory) · Manufacturing (Inventory & BOM)",
     "Warehouse + digital overlay; 99.8% accuracy vibe."),
 lib("LIB-14", "AI Finance, Billing & Invoicing",
     "Automated finance ops — invoices, GST/tax, reconciliation and payment flows on a clean finance dashboard; error-free, audit-ready.",
     "Finance (Expense / Receivable) · Hospital (Billing & Insurance Claims) · HRMS (Payroll & Compliance)",
     "Indian finance context (₹, GST) welcome."),
 lib("LIB-15", "AI Compliance, Security & Audit",
     "A glowing shield / check-badge over encrypted business data — compliance rules validated, audit trail logged, approvals cleared; enterprise-grade trust.",
     "Finance (Statutory Compliance) · Operations (Approval Engine) · HRMS (Payroll compliance)",
     "Shield + checkmark; secure, not scary."),
 lib("LIB-16", "AI HR, Recruitment & Onboarding",
     "AI streamlining hiring & onboarding — resume screening with fit-scores, a digital onboarding checklist, staff roster optimisation; a happy HR team.",
     "HRMS (Screening / Onboarding / Engagement) · Healthcare (Staff & Scheduling) · Hospital (Staff Scheduling)",
     "Diverse, professional, warm tone."),
 lib("LIB-17", "AI Logistics, Fleet & Route Optimization",
     "A logistics control room with an AI route-optimisation map — live vehicle pins, optimised routes, ETAs and delivery status; fuel/time saved.",
     "Supply Chain (Logistics Optimization) · Logistics (Route Optimization) · Logistics (Shipment Tracking) · Logistics (Freight Rate Automation)",
     "Map + fleet telemetry; India map acceptable."),
 lib("LIB-18", "AI Quality Vision & Inspection",
     "Computer-vision quality inspection on a production line — a camera detecting defects on products with bounding boxes and pass/fail marks at line speed.",
     "Manufacturing – Solution (Quality Control) · Manufacturing – Industry (Quality Defect Detection)",
     "99.4% accuracy; defect bounding boxes."),
 lib("LIB-19", "AI Predictive Maintenance & IoT",
     "IoT sensors + AI predicting machine failure before it happens — a machine with health gauges, vibration/temperature graphs and a 'predicted failure in X days' alert.",
     "Manufacturing – Solution (Predictive Maintenance) · Manufacturing – Industry (Predictive Maintenance)",
     "Industrial + sensor overlay; zero-downtime idea."),
 lib("LIB-20", "AI Personalization & Recommendations",
     "AI tailoring experiences per person — product/course/room recommendations adapting to an individual's behaviour; a delighted customer, relevant suggestions.",
     "Retail (Recommendations) · Hospitality (Guest Personalization) · Education (Learning Path)",
     "Show 'right thing, right person, right moment'."),
]

# ---------- HEROES ----------
def hero(id_, name, page, content, priority, reuse="Page-specific hero (single use).",
         notes="Full-bleed; left side darker for the headline; keep detail on the right half.",
         ratio="16:9", res="1920 x 1080 px (2x = 2560x1440 preferred)", orient="Landscape"):
    return dict(id=id_, name=name, where=f"{page} - Hero Section (full-bleed background)",
                purpose=f"Sets the tone for the {page}. First thing the visitor sees; must instantly signal the page topic + AI innovation.",
                content=content, style=HERO_STYLE, ratio=ratio, res=res,
                orient=orient, priority=priority, reuse=reuse, notes=notes)

HEROES = [
 hero("HERO-01","Sales Automation Hero","Solutions › Sales Automation",
      "A high-energy sales floor where an AI agent works alongside reps — a live pipeline/funnel hologram, deals moving to 'won', fast lead response.","High"),
 hero("HERO-02","Marketing Automation Hero","Solutions › Marketing Automation",
      "An AI marketing command centre pushing content & campaigns across channels — growth curve, content tiles and channel icons radiating out.","High"),
 hero("HERO-03","HRMS & Hiring Hero","Solutions › HRMS & Hiring",
      "A modern HR scene — AI screening candidates and running a digital onboarding flow; diverse team, hiring funnel and time-to-hire metric.","High"),
 hero("HERO-04","Operations Automation Hero","Solutions › Operations",
      "An operations nerve-centre where manual tasks are being automated — workflow board, document automation and approval flows on big screens.","High"),
 hero("HERO-05","Supply Chain Hero","Solutions › Supply Chain",
      "A connected supply-chain network — demand forecast, inventory and a logistics map linked as one intelligent flow from factory to doorstep.","High"),
 hero("HERO-06","Finance & Accounts Hero","Solutions › Finance & Accounts",
      "A CFO-grade finance control room — invoices, GST, reconciliation and cash-flow forecast automating on clean dashboards; audit-ready.","High"),
 hero("HERO-07","AI Voice & Chat Hero","Solutions › AI Voice & Chat",
      "An always-on AI agent handling voice + chat across WhatsApp, phone and web — headset/soundwave + chat bubbles, multilingual, 24/7.","High"),
 hero("HERO-08","Manufacturing (Solution) Hero","Solutions › Manufacturing",
      "A smart-factory floor (Industry 4.0) — machines with AI monitoring, quality-vision cameras and an OEE dashboard; zero unplanned downtime.","High",
      reuse="Distinct from HERO-15 (industry). This one is product/solution-led (dashboards & modules)."),
 hero("HERO-09","Hospital Management Hero","Solutions › Hospital Management",
      "A premium hospital operations view — OPD queue, billing and lab dashboards running; staff freed for patient care.","High"),
 hero("HERO-10","Healthcare Industry Hero","Industries › Healthcare",
      "A modern hospital lobby/OPD where doctors and staff use AI dashboards for appointments, records and billing — trust + digital health.","High"),
 hero("HERO-11","Education Industry Hero","Industries › Education",
      "A campus/edtech scene — admissions and student-success AI at work; students and faculty with a learning/analytics dashboard.","High"),
 hero("HERO-12","Hospitality Industry Hero","Industries › Hospitality",
      "A boutique hotel lobby/front desk with an AI guest & booking dashboard — warm, premium, guest-delight vibe.","High"),
 hero("HERO-13","Real Estate Industry Hero","Industries › Real Estate",
      "A property sales lounge with digital property listings and an AI lead-pipeline board — closing deals faster.","High"),
 hero("HERO-14","Retail & E-commerce Industry Hero","Industries › Retail & E-commerce",
      "An omnichannel retail scene — store + online + warehouse unified, AI shopping assistant and revenue-growth dashboard.","High"),
 hero("HERO-15","Manufacturing (Industry) Hero","Industries › Manufacturing",
      "A broad industrial plant view across sectors (auto, pharma, FMCG) with AI monitoring — factory-floor scale, sensors and safety.","High",
      reuse="Distinct from HERO-08 (solution). This one is sector/vertical-led (plant scale)."),
 hero("HERO-16","Logistics Industry Hero","Industries › Logistics",
      "A logistics control room + last-mile fleet — route-optimisation map, live tracking and on-time delivery; move smarter.","High"),
 hero("HERO-17","About Page Hero","About",
      "The AI Agentix team / company vision — a confident, diverse team in a modern office with subtle AI motifs; 'India's #1 AI automation partner'.","Medium",
      notes="Can lean more human/team than product. Left-aligned copy space."),
 hero("HERO-18","Solutions Index Hero","Solutions (index)",
      "A catalogue-style overview of AI solutions — modular tiles/agents representing every business function working as one system.","Medium"),
 hero("HERO-19","Industries Index Hero","Industries (index)",
      "A montage of industries served (healthcare, retail, logistics, etc.) unified under one AI platform — 'tuned to your industry'.","Medium"),
 hero("HERO-20","Technology Page Hero","Technology (planned)",
      "The AI/tech stack — LLMs, agents, integrations and infrastructure visualised as a premium, futuristic architecture.","Medium",
      notes="Future page — more abstract/futuristic allowed."),
 hero("HERO-21","Case Studies Page Hero","Case Studies (planned)",
      "A 'real results' montage — happy Indian business clients + rising metrics; proof and credibility.","Medium"),
 hero("HERO-22","Contact Page Hero / Banner","Contact (planned)",
      "An approachable 'let's talk' banner — a friendly team ready to help, subtle AI accents; inviting, low-pressure.","Medium",
      notes="Could be a wide banner (21:9) if layout needs — see notes.", ratio="16:9 (or 21:9 banner)"),
 hero("HERO-23","Blog / Insights Hero","Blog (planned)",
      "A knowledge/insights motif — articles, ideas and AI trends; editorial and clean.","Low"),
 hero("HERO-24","Careers Hero","Career (planned)",
      "A 'build the future of AI with us' scene — energetic, diverse team culture at AI Agentix.","Low"),
 hero("HERO-25","FAQ Hero / Banner","FAQ (planned)",
      "A calm 'questions answered' banner — supportive, clear, minimal.","Low"),
]

# ---------- CASE STUDIES ----------
def cs(id_, fname, industry, content):
    return dict(id=id_, name=f"Case Study — {industry}", where="Home › Case Studies (dark carousel) and Case Studies page",
                purpose=f"Visual for the {industry} success story card. Adds credibility/proof next to the metrics.",
                content=content, style=("Premium, moody/dark, cinematic — realistic industry scene with a subtle AI/"
                                        "dashboard element and orange accent. Same style family as the rest of the site."),
                ratio="4:3", res="1600 × 1200 px (min 1280×960)", orient="Landscape", priority="Medium",
                reuse="Single use (one per case study).",
                notes=f"Deliver as `{fname}`. Keep the main subject centre-left (a small industry tag chip sits bottom-right). "
                      f"Dark/cinematic so white overlay text stays readable.")

CASE_STUDIES = [
 cs("CS-01","cs-real-estate","Real Estate (PropTech, Mumbai)","Modern real-estate sales office / agent with an AI lead-pipeline dashboard on screen."),
 cs("CS-02","cs-ecommerce","E-commerce (D2C, Bengaluru)","A D2C brand fulfilment/warehouse or store with an AI support-chat dashboard overlay."),
 cs("CS-03","cs-healthcare","Healthcare (Clinic Chain, New Delhi)","A clinic reception/OPD with an AI appointment & reminder dashboard."),
 cs("CS-04","cs-education","Education (CBSE School, Pune)","A school admissions desk / campus with an AI enrolment dashboard."),
 cs("CS-05","cs-logistics","Logistics (Last-Mile, Delhi NCR)","A last-mile dispatch/control room with a route-optimisation map dashboard."),
 cs("CS-06","cs-hospitality","Hospitality (Boutique Hotel, Rajasthan)","A boutique hotel lobby / front desk with an AI booking & guest dashboard."),
 cs("CS-07","cs-manufacturing","Manufacturing (Auto Components, Pune)","An auto-components factory floor with an AI machine-monitoring / OEE dashboard."),
]

# ---------- MISC ----------
MISC = [
 dict(id="MISC-01", name="About — Mission Image", where="About › Mission section (portrait card, hover reveals 2026 targets)",
      purpose="Anchors the mission section; humanises the brand.",
      content="A purposeful AI-in-business moment — a team or founder with a subtle AI dashboard, conveying vision & momentum.",
      style=LIB_STYLE, ratio="4:5", res="1080 × 1350 px", orient="Portrait", priority="Medium",
      reuse="Single use.", notes="Portrait; hover overlays a dark panel, so a slightly darker image works well."),
 dict(id="MISC-02", name="About — Team Photos (real)", where="About › Team section (grid)",
      purpose="Real credibility — actual team members.",
      content="Real headshots/portraits of the actual AI Agentix team on a consistent background.",
      style="Consistent, professional headshots (real people — NOT AI-generated/stock).",
      ratio="1:1", res="800 × 800 px each", orient="Square", priority="Low",
      reuse="One per team member.",
      notes="These must be REAL team photos supplied by the client, shot on a consistent neutral/branded background. Not a designer illustration."),
]

# =========================================================================
# BUILD DOCUMENT
# =========================================================================
doc = Document()
# base styles
normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(10.5)
for i, sz in [(1, 18), (2, 14), (3, 11.5)]:
    st = doc.styles[f'Heading {i}']; st.font.color.rgb = RGBColor.from_string(NAVY if i == 1 else DARK); st.font.size = Pt(sz)

# ---- COVER PAGE ----
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("AI AGENTIX"); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = RGBColor.from_string(ORANGE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Master Image Asset Requirement Document"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Designer Handoff Document"); r.font.size = Pt(14); r.italic = True; r.font.color.rgb = RGBColor.from_string(DARK)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Website: aiagentix.com  ·  Brand palette: Orange #F26522 + Blue #2F80ED  ·  Style: Photoreal AI-in-business")
r.font.size = Pt(11)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Prepared for the Graphic Design Team"); r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string("666666")

# ---- HOW TO USE / INTRO ----
doc.add_page_break()
doc.add_heading("How to Use This Document", level=1)
intro = [
 ("Goal", "This document lists every image the AI Agentix website needs. The Home Page is finalised and already uses "
          "approved images (from the earlier ZIP) — those are NOT re-requested here. Everything else is a 'Pending "
          "Designer Asset' captured below."),
 ("Small, reusable set (important)", "We deliberately did NOT create ~110 one-off images. Instead there is a compact "
          "REUSABLE CONTENT LIBRARY of 20 images (Section 2). Each library image is used across MANY pages (2–7 places), "
          "so the design effort — and the site's visual consistency — stays high while the number of assets stays low. "
          "Page Heroes (Section 3), Case Studies (Section 4) and a few section images (Section 5) are the only "
          "single-use assets."),
 ("Brand consistency", COLORS),
 ("Consistent style", "Every image must match the approved Home Page style: photorealistic 'AI-in-business' scenes with "
          "subtle holographic UI / dashboard elements, premium and enterprise-grade."),
 ("Delivery", "Deliver source as PNG or JPG at the stated resolution (we convert to optimised WebP). Use the exact "
          "'Filename' where given. Keep the main subject centre-left where noted (headlines/overlays sit on the left)."),
 ("Home Hero = video", "The Home Page hero uses a VIDEO (Section 7), NOT an image — do not create a Home hero image."),
]
for k, v in intro:
    p = doc.add_paragraph(); r = p.add_run(f"{k}:  "); r.bold = True; r.font.color.rgb = RGBColor.from_string(ORANGE)
    p.add_run(v)
add_toc(doc)

# ---- SECTION 1: TOC placeholder heading already; now sections ----
section_title(doc, 1, "Reusable Content Image Library (20 images)",
              "These 20 images cover ALL Solution-capability panels and Industry-solution rows across the site. Each is "
              "reused on multiple pages (see 'Reuse / Used On'). No two of these appear twice on the same page — reuse "
              "happens across different pages. This is the core of the 'small, reusable set' approach.")
for e in LIBRARY:
    image_entry(doc, e)

section_title(doc, 2, "Page Hero Images (per page)",
              "One cinematic 16:9 hero per page (except the Home Page, which uses a video). Solution & Industry heroes are "
              "High priority (those pages are live). Index/other/future-page heroes are Medium/Low.")
for e in HEROES:
    image_entry(doc, e)

section_title(doc, 3, "Case Study Images (7 images)",
              "Dark, cinematic proof images for the Case Studies carousel (Home) and the Case Studies page. One per story.")
for e in CASE_STUDIES:
    image_entry(doc, e)

section_title(doc, 4, "Other / Section Images",
              "A few remaining single-use section images.")
for e in MISC:
    image_entry(doc, e)

# ---- SECTION 5: HOME HERO VIDEO ----
doc.add_page_break()
doc.add_heading("5. Home Page Hero Video Requirement", level=1)
vid = [
 ("Purpose", "The Home Page hero is the brand's first impression. A short, seamless, looping background video conveys "
             "'AI automation running your business 24/7' with more energy and premium feel than a static image. Website "
             "headline + CTA text sit ON TOP of the video (kept editable in code), so the video must NOT contain text."),
 ("Complete Video Concept", "A cinematic journey through an 'AI-powered business': we glide through a modern enterprise "
             "where autonomous AI agents quietly run sales, support, finance, operations and logistics — holographic "
             "dashboards, flowing data, and calm, in-control humans. The mood is premium, futuristic-but-real, optimistic. "
             "Orange energy pulses trace the automation as it moves from task to task."),
 ("Scene Breakdown", "Scene 1 (0–3s): Wide push-in through a sleek office; subtle orange data streams ignite. "
             "Scene 2 (3–6s): Close on an AI agent/hologram qualifying leads & sending WhatsApp/email — chat bubbles fly. "
             "Scene 3 (6–9s): A live analytics/BI dashboard assembles itself; KPIs tick up. "
             "Scene 4 (9–12s): Pull back to reveal the whole business connected as one glowing network (sales→ops→finance→"
             "logistics). Scene 4 dissolves seamlessly back into Scene 1 for the loop."),
 ("Camera Movement", "Smooth, slow, continuous motion — gentle push-ins, glides and parallax. No hard cuts; use graceful "
             "dissolves. Steadicam/drone feel, never shaky."),
 ("Motion Style", "Cinematic + subtle motion-graphics. Real environments with holographic UI overlays and flowing "
             "orange/blue data particles. Elegant, not busy."),
 ("Lighting", "Dark, moody, high-end; soft key light on subjects, orange rim-light accents, cool blue ambient for the "
             "tech/holographic elements."),
 ("Mood", "Premium · innovative · trustworthy · calm-in-control · optimistic."),
 ("Colour Palette", "Brand Orange #F26522 as the hero energy colour, secondary Blue #2F80ED / deep navy #0B1F3A for "
             "tech/holographic elements, over dark charcoal #12141A bases. No off-brand colours."),
 ("Duration", "10–15 seconds (seamless loop). 12s recommended."),
 ("Loop Behaviour", "MUST loop seamlessly — the last frame should blend invisibly into the first (design Scene 4 → Scene 1 "
             "as a match-dissolve). No visible jump/flash."),
 ("Resolution", "3840 × 2160 (4K) master; export a 1920 × 1080 (1080p) optimised web version."),
 ("Aspect Ratio", "16:9 (primary). Also provide a 9:16 or centre-safe crop guidance for mobile, OR keep key action "
             "centred so a CSS cover-crop works on mobile."),
 ("Frame Rate", "24 fps (cinematic) or 30 fps. Keep consistent."),
 ("Output Formats", "MP4 (H.264) for broad support + WebM (VP9) for smaller size. Provide both; plus a compressed "
             "web-optimised MP4 (target < 6–8 MB for the 1080p loop)."),
 ("Compression Recommendation", "Two-pass encode; strip audio track; optimise for silent autoplay background use. Target "
             "the loop under ~8 MB at 1080p without visible banding (watch dark gradients)."),
 ("Background Music", "Not required (plays muted/autoplay as a background loop). If a version with sound is ever needed, "
             "keep it subtle, modern and royalty-free — but the website will play it MUTED."),
 ("Text in Video", "NO embedded text. All headline/subtext/CTA remain editable HTML on the website, layered over the video."),
 ("Designer / Production Notes", "Keep the centre-left area visually calmer (the headline sits there). Ensure strong "
             "contrast so white text is readable over the video at all times (a subtle dark gradient overlay will be added "
             "in code). Deliver the seamless loop + a poster/first-frame still (1920×1080) for fast initial paint."),
]
t = doc.add_table(rows=0, cols=2); t.style = 'Table Grid'
for k, v in vid:
    c = t.add_row().cells
    set_text(c[0], k, bold=True, size=9.5, color=DARK); shade(c[0], LIGHT)
    set_text(c[1], v, size=9.5)
    c[0].width = Inches(1.9); c[1].width = Inches(4.7)

# ---- SECTION 6: SUMMARY TABLES ----
doc.add_page_break()
doc.add_heading("6. Final Summary Tables", level=1)

ALL = LIBRARY + HEROES + CASE_STUDIES + MISC

def simple_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for i, h in enumerate(headers):
        set_text(t.rows[0].cells[i], h, bold=True, size=9, color=WHITE); shade(t.rows[0].cells[i], ORANGE)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_text(cells[i], str(val), size=9)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    return t

# 6.1 Inventory
doc.add_heading("6.1  Image Inventory", level=2)
inv_rows = [(e['id'], e['name'], e['where'].split(' · ')[0].split(' › ')[0] if ' · ' in e['where'] else e['where'][:40],
             e.get('priority', '—')) for e in ALL]
simple_table(doc, ["Image ID", "Image Name", "Page / Area", "Priority"],
             [(e['id'], e['name'], (e['where'][:46] + '…') if len(e['where']) > 47 else e['where'], e['priority']) for e in ALL],
             widths=[0.9, 2.4, 2.6, 0.8])

# 6.2 Aspect ratio summary
def count_by(key_fn):
    d = {}
    for e in ALL:
        k = key_fn(e); d[k] = d.get(k, 0) + 1
    return sorted(d.items(), key=lambda x: -x[1])

doc.add_heading("6.2  Aspect Ratio Summary", level=2)
simple_table(doc, ["Aspect Ratio", "Number of Images"], count_by(lambda e: e['ratio']), widths=[2.5, 2.0])

doc.add_heading("6.3  Resolution Summary", level=2)
simple_table(doc, ["Recommended Resolution", "Number of Images"], count_by(lambda e: e['res']), widths=[3.4, 2.0])

doc.add_heading("6.4  Orientation Summary", level=2)
simple_table(doc, ["Orientation", "Number of Images"], count_by(lambda e: e['orient']), widths=[2.5, 2.0])

doc.add_heading("6.5  File Format Summary", level=2)
simple_table(doc, ["File Format", "Number of Images"], count_by(lambda e: e.get('fmt', 'PNG or JPG (→ WebP)')), widths=[4.2, 1.6])

doc.add_heading("6.6  Priority Breakdown", level=2)
simple_table(doc, ["Priority", "Number of Images"], count_by(lambda e: e['priority']), widths=[2.5, 2.0])

# 6.7 totals
doc.add_heading("6.7  Totals", level=2)
n_lib = len(LIBRARY); n_hero = len(HEROES); n_cs = len(CASE_STUDIES); n_misc = len(MISC)
tot_rows = [
 ("Reusable content library images", n_lib, "Cover 90+ capability/solution placements across the site (heavy reuse)"),
 ("Page hero images", n_hero, "One per page; Home uses a video instead"),
 ("Case study images", n_cs, "One per success story"),
 ("Other / section images", n_misc, "Incl. real team photos (client-supplied)"),
 ("TOTAL UNIQUE IMAGES TO CREATE", n_lib + n_hero + n_cs + n_misc, "vs ~110 if every slot were unique — the library cuts this dramatically"),
 ("Home Page Hero Video", 1, "Separate deliverable (Section 5)"),
]
simple_table(doc, ["Asset Group", "Count", "Notes"], tot_rows, widths=[2.7, 0.7, 3.0])

hi = [e for e in ALL if e['priority'] == 'High']; me = [e for e in ALL if e['priority'] == 'Medium']; lo = [e for e in ALL if e['priority'] == 'Low']
doc.add_heading("6.8  High Priority Images (do first)", level=2)
simple_table(doc, ["Image ID", "Image Name"], [(e['id'], e['name']) for e in hi], widths=[1.0, 4.5])
doc.add_heading("6.9  Medium Priority Images", level=2)
simple_table(doc, ["Image ID", "Image Name"], [(e['id'], e['name']) for e in me], widths=[1.0, 4.5])
doc.add_heading("6.10  Low Priority Images", level=2)
simple_table(doc, ["Image ID", "Image Name"], [(e['id'], e['name']) for e in lo], widths=[1.0, 4.5])

# footer note
doc.add_paragraph()
p = doc.add_paragraph(); r = p.add_run("End of document — prepared for the AI Agentix design team.")
r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string("777777")

OUT = r"D:\AI AGENTIX\AI Agentix - Image Asset Requirement Document.docx"
doc.save(OUT)
print("SAVED:", OUT)
print("Totals -> library:", n_lib, "heroes:", n_hero, "case_studies:", n_cs, "misc:", n_misc,
      "TOTAL:", n_lib + n_hero + n_cs + n_misc)
